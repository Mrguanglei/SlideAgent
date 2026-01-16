"""
知识库文档解析器

支持的文件格式：
- PDF (使用 PyPDF2)
- Word (docx, 使用 python-docx)
- Excel (xlsx, 使用 openpyxl)
- HTML (使用 BeautifulSoup)
- TXT, Markdown (直接读取)
- URL (使用 requests + BeautifulSoup)
"""

import os
import re
import chardet
from typing import Optional, Dict, Any, Tuple
from pathlib import Path

# PDF 解析
try:
    from PyPDF2 import PdfReader
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

# Word 解析
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Excel 解析
try:
    from openpyxl import load_workbook
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

# HTML 解析
from bs4 import BeautifulSoup
import requests


class DocumentParser:
    """文档解析器 - 从各种格式中提取文本"""
    
    # 支持的文件扩展名
    SUPPORTED_EXTENSIONS = {
        'pdf': 'PDF 文档',
        'docx': 'Word 文档',
        'doc': 'Word 文档 (旧版)',
        'xlsx': 'Excel 表格',
        'xls': 'Excel 表格 (旧版)',
        'txt': '纯文本',
        'md': 'Markdown',
        'html': 'HTML 网页',
        'htm': 'HTML 网页',
        'xml': 'XML 文档',
    }
    
    @classmethod
    def get_file_type(cls, filename: str) -> Optional[str]:
        """获取文件类型"""
        ext = Path(filename).suffix.lower().lstrip('.')
        return ext if ext in cls.SUPPORTED_EXTENSIONS else None
    
    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """检查文件是否支持"""
        return cls.get_file_type(filename) is not None
    
    @classmethod
    async def parse(cls, file_path: str, file_type: Optional[str] = None) -> Tuple[str, Dict[str, Any]]:
        """
        解析文档并提取文本
        
        Args:
            file_path: 文件路径
            file_type: 文件类型（可选，自动检测）
            
        Returns:
            (提取的文本内容, 元数据字典)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        if file_type is None:
            file_type = cls.get_file_type(file_path)
        
        if file_type is None:
            raise ValueError(f"不支持的文件格式: {file_path}")
        
        # 根据文件类型选择解析方法
        parser_map = {
            'pdf': cls._parse_pdf,
            'docx': cls._parse_docx,
            'doc': cls._parse_docx,  # 尝试用 docx 解析
            'xlsx': cls._parse_xlsx,
            'xls': cls._parse_xlsx,  # 尝试用 openpyxl 解析
            'txt': cls._parse_text,
            'md': cls._parse_text,
            'html': cls._parse_html,
            'htm': cls._parse_html,
            'xml': cls._parse_xml,
        }
        
        parser = parser_map.get(file_type)
        if parser is None:
            raise ValueError(f"不支持的文件类型: {file_type}")
        
        return await parser(file_path)
    
    @classmethod
    async def parse_url(cls, url: str) -> Tuple[str, Dict[str, Any]]:
        """
        解析网页 URL 并提取文本
        
        Args:
            url: 网页 URL
            
        Returns:
            (提取的文本内容, 元数据字典)
        """
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # 检测编码
            if response.encoding == 'ISO-8859-1':
                detected = chardet.detect(response.content)
                response.encoding = detected.get('encoding', 'utf-8')
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # 移除脚本和样式
            for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
                tag.decompose()
            
            # 获取标题
            title = soup.title.string if soup.title else ''
            
            # 提取正文
            text = soup.get_text(separator='\n', strip=True)
            text = cls._clean_text(text)
            
            metadata = {
                'title': title,
                'url': url,
                'content_length': len(text),
            }
            
            return text, metadata
            
        except Exception as e:
            raise ValueError(f"解析 URL 失败: {str(e)}")
    
    @classmethod
    async def parse_text_content(cls, content: str) -> Tuple[str, Dict[str, Any]]:
        """
        解析纯文本内容
        
        Args:
            content: 文本内容
            
        Returns:
            (清洗后的文本, 元数据字典)
        """
        text = cls._clean_text(content)
        metadata = {
            'content_length': len(text),
            'source': 'text_input',
        }
        return text, metadata
    
    # ==================== 私有解析方法 ====================
    
    @classmethod
    async def _parse_pdf(cls, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析 PDF 文件"""
        if not HAS_PYPDF2:
            raise ImportError("PyPDF2 未安装，无法解析 PDF")
        
        try:
            reader = PdfReader(file_path)
            text_parts = []
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text() or ''
                if page_text.strip():
                    text_parts.append(f"[第 {page_num + 1} 页]\n{page_text}")
            
            text = '\n\n'.join(text_parts)
            text = cls._clean_text(text)
            
            metadata = {
                'page_count': len(reader.pages),
                'content_length': len(text),
            }
            
            # 尝试获取文档信息
            if reader.metadata:
                if reader.metadata.title:
                    metadata['title'] = reader.metadata.title
                if reader.metadata.author:
                    metadata['author'] = reader.metadata.author
            
            return text, metadata
            
        except Exception as e:
            raise ValueError(f"PDF 解析失败: {str(e)}")
    
    @classmethod
    async def _parse_docx(cls, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析 Word 文档"""
        if not HAS_DOCX:
            raise ImportError("python-docx 未安装，无法解析 Word 文档")
        
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            text = '\n\n'.join(text_parts)
            text = cls._clean_text(text)
            
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'content_length': len(text),
            }
            
            # 尝试获取文档属性
            if doc.core_properties:
                if doc.core_properties.title:
                    metadata['title'] = doc.core_properties.title
                if doc.core_properties.author:
                    metadata['author'] = doc.core_properties.author
            
            return text, metadata
            
        except Exception as e:
            raise ValueError(f"Word 文档解析失败: {str(e)}")
    
    @classmethod
    async def _parse_xlsx(cls, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析 Excel 文件"""
        if not HAS_OPENPYXL:
            raise ImportError("openpyxl 未安装，无法解析 Excel 文件")
        
        try:
            wb = load_workbook(file_path, read_only=True, data_only=True)
            text_parts = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_text = [f"[工作表: {sheet_name}]"]
                
                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else '' for cell in row]
                    row_text = ' | '.join(v for v in row_values if v)
                    if row_text:
                        sheet_text.append(row_text)
                
                if len(sheet_text) > 1:
                    text_parts.append('\n'.join(sheet_text))
            
            wb.close()
            
            text = '\n\n'.join(text_parts)
            text = cls._clean_text(text)
            
            metadata = {
                'sheet_count': len(wb.sheetnames),
                'sheet_names': wb.sheetnames,
                'content_length': len(text),
            }
            
            return text, metadata
            
        except Exception as e:
            raise ValueError(f"Excel 文件解析失败: {str(e)}")
    
    @classmethod
    async def _parse_text(cls, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析纯文本文件"""
        try:
            # 检测编码
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                detected = chardet.detect(raw_data)
                encoding = detected.get('encoding', 'utf-8')
            
            # 读取文本
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                text = f.read()
            
            text = cls._clean_text(text)
            
            metadata = {
                'encoding': encoding,
                'content_length': len(text),
            }
            
            return text, metadata
            
        except Exception as e:
            raise ValueError(f"文本文件解析失败: {str(e)}")
    
    @classmethod
    async def _parse_html(cls, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析 HTML 文件"""
        try:
            # 检测编码
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                detected = chardet.detect(raw_data)
                encoding = detected.get('encoding', 'utf-8')
            
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 移除脚本和样式
            for tag in soup(['script', 'style']):
                tag.decompose()
            
            # 获取标题
            title = soup.title.string if soup.title else ''
            
            # 提取文本
            text = soup.get_text(separator='\n', strip=True)
            text = cls._clean_text(text)
            
            metadata = {
                'title': title,
                'content_length': len(text),
            }
            
            return text, metadata
            
        except Exception as e:
            raise ValueError(f"HTML 文件解析失败: {str(e)}")
    
    @classmethod
    async def _parse_xml(cls, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析 XML 文件"""
        try:
            # 检测编码
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                detected = chardet.detect(raw_data)
                encoding = detected.get('encoding', 'utf-8')
            
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                xml_content = f.read()
            
            soup = BeautifulSoup(xml_content, 'xml')
            text = soup.get_text(separator='\n', strip=True)
            text = cls._clean_text(text)
            
            metadata = {
                'content_length': len(text),
            }
            
            return text, metadata
            
        except Exception as e:
            raise ValueError(f"XML 文件解析失败: {str(e)}")
    
    # ==================== 文本清洗 ====================
    
    @classmethod
    def _clean_text(cls, text: str) -> str:
        """清洗文本内容"""
        if not text:
            return ''
        
        # 替换多个空白字符为单个空格
        text = re.sub(r'[ \t]+', ' ', text)
        
        # 替换多个换行为两个换行
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # 去除每行首尾空白
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # 去除首尾空白
        text = text.strip()
        
        return text
